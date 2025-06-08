import unittest
import os
import shutil
from datetime import datetime
from parse_folder_path import parse_folder_path
from pack_search import pack_search
from docx import Document
import pandas as pd

class TestParseFolderPath(unittest.TestCase):
    def setUp(self):
        # 设置测试目录路径
        self.test_dir = r"D:\2023级计科2班(备份日期2025年5月1日)"
        
        # 检查测试目录是否存在
        if not os.path.exists(self.test_dir):
            raise FileNotFoundError(f"测试目录不存在: {self.test_dir}")
    
    def test_parse_folder_path(self):
        try:
            # 调用函数
            result = parse_folder_path(self.test_dir)
            
            # 验证返回结果是否为列表
            self.assertIsInstance(result, list, "返回值不是列表类型")
            
            # 如果目录中有文件，验证列表内容
            if result:
                # 验证列表中的每个元素都是字典且包含必要的键
                for file_info in result:
                    self.assertIsInstance(file_info, dict, "文件信息不是字典类型")
                    required_keys = ["name", "absolute_path", "extension", 
                                   "created_time", "size_bytes", "content"]
                    for key in required_keys:
                        self.assertIn(key, file_info, f"文件信息缺少必要键: {key}")
            
            print("测试成功")
            return True
        except Exception as e:
            print(f"测试失败: {str(e)}")
            return False

class TestFileParsing(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        # 创建测试目录和测试文件
        cls.test_dir = "test_files"
        os.makedirs(cls.test_dir, exist_ok=True)
        
        # 创建测试PDF文件
        cls.pdf_path = os.path.join(cls.test_dir, "testPDF.pdf")
        if not os.path.exists(cls.pdf_path):
            with open(cls.pdf_path, 'wb') as f:
                f.write(b'%PDF-1.4\n%' + b' ' * 100 + b'\n%%EOF')  # 创建一个简单的PDF文件
        
        # 创建测试Word文件
        cls.docx_path = os.path.join(cls.test_dir, "testDOC.docx")
        if not os.path.exists(cls.docx_path):
            doc = Document()
            doc.add_paragraph("这是一个测试Word文档")
            doc.add_paragraph("包含多行文本")
            doc.save(cls.docx_path)
        
        # 创建测试Excel文件
        cls.xlsx_path = os.path.join(cls.test_dir, "testXLSX.xlsx")
        if not os.path.exists(cls.xlsx_path):
            df = pd.DataFrame({
                'Name': ['Alice', 'Bob', 'Charlie'],
                'Age': [25, 30, 35],
                'City': ['New York', 'London', 'Paris']
            })
            df.to_excel(cls.xlsx_path, index=False)
    
    def test_pdf_file(self):
        """测试PDF文件解析"""
        result = parse_folder_path(self.test_dir)
        pdf_files = [f for f in result if f['extension'] == '.pdf']
        
        self.assertGreater(len(pdf_files), 0, "没有找到PDF文件")
        pdf_info = pdf_files[0]
        
        print("\nPDF文件测试结果:")
        print(f"文件名: {pdf_info['name']}")
        print(f"路径: {pdf_info['absolute_path']}")
        print(f"大小: {pdf_info['size_bytes']} 字节")
        print(f"创建时间: {pdf_info['created_time']}")
        print(f"内容预览: {pdf_info['content'][:100]}...")
        
        self.assertTrue(isinstance(pdf_info['content'], str))
    
    def test_word_file(self):
        """测试Word文件解析"""
        result = parse_folder_path(self.test_dir)
        docx_files = [f for f in result if f['extension'] == '.docx']
        
        self.assertGreater(len(docx_files), 0, "没有找到Word文件")
        docx_info = docx_files[0]
        
        print("\nWord文件测试结果:")
        print(f"文件名: {docx_info['name']}")
        print(f"路径: {docx_info['absolute_path']}")
        print(f"大小: {docx_info['size_bytes']} 字节")
        print(f"创建时间: {docx_info['created_time']}")
        print(f"内容: {docx_info['content']}")
        
        self.assertIn("这是一个测试Word文档", docx_info['content'])
        self.assertIn("包含多行文本", docx_info['content'])
    
    def test_excel_file(self):
        """测试Excel文件解析"""
        result = parse_folder_path(self.test_dir)
        xlsx_files = [f for f in result if f['extension'] == '.xlsx']
        
        self.assertGreater(len(xlsx_files), 0, "没有找到Excel文件")
        xlsx_info = xlsx_files[0]
        
        print("\nExcel文件测试结果:")
        print(f"文件名: {xlsx_info['name']}")
        print(f"路径: {xlsx_info['absolute_path']}")
        print(f"大小: {xlsx_info['size_bytes']} 字节")
        print(f"创建时间: {xlsx_info['created_time']}")
        print(f"内容预览: {xlsx_info['content'][:100]}...")
        
        self.assertIn("Alice", xlsx_info['content'])
        self.assertIn("Bob", xlsx_info['content'])
        self.assertIn("Charlie", xlsx_info['content'])
    
    @classmethod
    def tearDownClass(cls):
        # 清理测试文件
        if os.path.exists(cls.test_dir):
            shutil.rmtree(cls.test_dir)

def paser_file(file_path):
    """解析单个文件并返回新路径和解析结果"""
    directory = os.path.dirname(file_path)
    filename = os.path.basename(file_path)
    
    # 使用parse_folder_path函数解析
    results = parse_folder_path(directory)
    
    # 查找匹配的文件
    for file_info in results:
        if file_info['name'] == filename:
            new_path = os.path.join("parsed_files", filename)
            reason = "解析成功" if not file_info['content'].startswith("<") else file_info['content']
            
            # 确保输出目录存在
            os.makedirs("parsed_files", exist_ok=True)
            
            return {
                "newPath": new_path,
                "reason": reason,
                "fileInfo": file_info
            }
    
    return {
        "newPath": "",
        "reason": "文件未找到",
        "fileInfo": None
    }

def test_search_function():
    """测试搜索功能"""
    query = "关于奖学金的文件"
    print("\n测试搜索功能:")
    print(f"搜索查询: '{query}'")
    results = pack_search(query)
    print("搜索结果:", results)
    return results

if __name__ == "__main__":
    # 运行文件解析测试
    print("="*50)
    print("开始运行文件解析测试")
    print("="*50)
    
    # 运行目录解析测试
    print("\n运行目录解析测试...")
    tester = TestParseFolderPath()
    try:
        tester.setUp()
        test_result = tester.test_parse_folder_path()
        
        if test_result:
            print("测试结果: 成功")
        else:
            print("测试结果: 失败")
    except FileNotFoundError as e:
        print(f"测试初始化失败: {str(e)}")
        print("测试结果: 失败 - 目录不存在")
    except Exception as e:
        print(f"测试过程中发生意外错误: {str(e)}")
        print("测试结果: 失败 - 意外错误")
    
    # 运行文件类型测试
    print("\n" + "="*50)
    print("开始运行文件类型测试")
    print("="*50)
    
    # 创建测试文件目录
    os.makedirs("test_files", exist_ok=True)
    
    # 测试PDF文件
    pdf_result = paser_file("test_files/testPDF.pdf")
    print("\nPDF文件解析结果:")
    print(f"新路径: {pdf_result['newPath']}")
    print(f"解析结果: {pdf_result['reason']}")
    
    # 测试Word文件
    docx_result = paser_file("test_files/testDOC.docx")
    print("\nWord文件解析结果:")
    print(f"新路径: {docx_result['newPath']}")
    print(f"解析结果: {docx_result['reason']}")
    
    # 测试Excel文件
    xlsx_result = paser_file("test_files/testXLSX.xlsx")
    print("\nExcel文件解析结果:")
    print(f"新路径: {xlsx_result['newPath']}")
    print(f"解析结果: {xlsx_result['reason']}")
    
    # 运行单元测试
    print("\n" + "="*50)
    print("开始运行单元测试")
    print("="*50)
    unittest.main(argv=['first-arg-is-ignored'], exit=False)
    
    # 测试搜索功能
    print("\n" + "="*50)
    print("开始运行搜索功能测试")
    print("="*50)
    search_results = test_search_function()
    
    print("\n" + "="*50)
    print("所有测试完成")
    print("="*50)